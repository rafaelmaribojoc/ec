import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_form_details(filepath):
    """Extract detailed form structure from a docx file"""
    print(f"\n{'='*80}")
    print(f"FILE: {os.path.basename(filepath)}")
    print(f"{'='*80}\n")
    
    try:
        doc = Document(filepath)
        
        # Document properties
        print(f"[DOCUMENT INFO]")
        print(f"Sections: {len(doc.sections)}")
        print(f"Paragraphs: {len(doc.paragraphs)}")
        print(f"Tables: {len(doc.tables)}")
        print()
        
        # Extract paragraphs with styling
        print("[CONTENT]")
        print("-" * 40)
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Get alignment
                alignment = "LEFT"
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    alignment = "CENTER"
                elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    alignment = "RIGHT"
                elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    alignment = "JUSTIFY"
                
                # Get style info
                style_name = para.style.name if para.style else "Normal"
                
                # Check if bold
                is_bold = False
                if para.runs:
                    is_bold = any(run.bold for run in para.runs if run.bold is not None)
                
                # Print with formatting hints
                fmt_hints = []
                if is_bold:
                    fmt_hints.append("BOLD")
                if alignment != "LEFT":
                    fmt_hints.append(alignment)
                if style_name != "Normal":
                    fmt_hints.append(f"Style:{style_name}")
                
                hint_str = f" [{', '.join(fmt_hints)}]" if fmt_hints else ""
                print(f"{text}{hint_str}")
        
        # Extract tables
        if doc.tables:
            print(f"\n[TABLES: {len(doc.tables)}]")
            print("-" * 40)
            
            for idx, table in enumerate(doc.tables):
                print(f"\n--- Table {idx + 1} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
                for row_idx, row in enumerate(table.rows):
                    cells = [cell.text.strip().replace('\n', ' | ') for cell in row.cells]
                    # Remove duplicates from merged cells
                    unique_cells = []
                    prev = None
                    for c in cells:
                        if c != prev:
                            unique_cells.append(c)
                            prev = c
                    print(f"  Row {row_idx}: {' | '.join(unique_cells)}")
                    
                    if row_idx > 25:  # Limit rows shown
                        print(f"  ... ({len(table.rows) - 25} more rows)")
                        break
        
        print()
        
    except Exception as e:
        print(f"Error processing {filepath}: {e}")

# All form templates
forms = [
    ("Social Service", [
        "form_templates/Social Service/2025 CASE FOLDER.docx",
    ]),
    ("Home Life Service", [
        "form_templates/Home Life Service/FINAL INVENTORY UPON DISCHARGE 2025.docx",
        "form_templates/Home Life Service/FINAL PROGRESS NOTES 2025.docx",
        "form_templates/Home Life Service/FINAL INVENTORY UPON ADMISSION 2025.docx",
        "form_templates/Home Life Service/FINAL INCIDENT REPORT 2025.docx",
        "form_templates/Home Life Service/FIINAL NEW OUT ON PASS 1.docx",
        "form_templates/Home Life Service/FINAL INVENTORY REPORTS 2025.docx",
    ]),
    ("Psychological Service", [
        "form_templates/Psychological Service/Psych Service Progress Notes.docx",
        "form_templates/Psychological Service/Psych Service Group Session I Activity.docx",
        "form_templates/Psychological Service/Inter-Service Referral (1).docx",
        "form_templates/Psychological Service/Individual Sessions Report Blank Template.docx",
        "form_templates/Psychological Service/Initial Psychological Assessment.docx",
        "form_templates/Psychological Service/Psychometricians Report.docx",
    ]),
]

base_path = "E:/Capstone/ElderCare"

for service_name, files in forms:
    print(f"\n{'#'*80}")
    print(f"# {service_name.upper()}")
    print(f"{'#'*80}")
    
    for filepath in files:
        full_path = os.path.join(base_path, filepath)
        if os.path.exists(full_path):
            extract_form_details(full_path)
        else:
            print(f"File not found: {filepath}")
